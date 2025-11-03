import os
import fitz  # PyMuPDF
from docx import Document
import re

class ResumeParser:
    def parse(self, file_path):
        """Parse resume and extract text"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_pdf(self, file_path):
        """Extract text from PDF using PyMuPDF"""
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            if not text.strip():
                raise ValueError("PDF appears to be empty or unreadable")
            
            return text
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    def _parse_docx(self, file_path):
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            if not text.strip():
                raise ValueError("DOCX appears to be empty")
            
            return text
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.,;:()\-@#+]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()